#!/usr/bin/env python3
"""Fill AInIS checklists from the living IJIET A29/A30 pack. Honest Đạt/Chưa."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Twips

HERE = Path(__file__).resolve().parent
BACKUP = HERE / "_blank_backup_20260901"
BACKUP.mkdir(exist_ok=True)

TITLE = (
    "Reproducible Sparse-Concept and Calibration Diagnostics for Knowledge Tracing"
)
AUTHORS_VN = (
    "Nguyễn Khánh Trình, Đào Minh Tuấn, Nguyễn Tiến Dương, "
    "Nguyễn Chí Thành, Nguyễn Văn Hậu"
)
AUTHORS_EN = (
    "Khanh-Trinh Nguyen, Tuan Dao Minh, Duong Nguyen Tien, "
    "Chi Thanh Nguyen, Van-Hau Nguyen"
)
NCS = "Nguyễn Khánh Trình"
VENUE = "IJIET (International Journal of Information and Education Technology, www.ijiet.org)"
DATE_VN = "01/09/2026"
DATE_LINE = "Hưng Yên, ngày 01 tháng 09 năm 2026"

# Proposed backups for GS — not treated as already approved.
BACKUP1 = "IEEE Transactions on Learning Technologies (đề xuất, chờ GS chốt)"
BACKUP2 = "International Journal of Artificial Intelligence in Education (đề xuất, chờ GS chốt)"


def set_run_text(para, text: str) -> None:
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def fill_after_label(para, new_value: str) -> None:
    """Keep the first run (label), replace remaining runs with the value."""
    if len(para.runs) >= 2:
        para.runs[1].text = new_value
        for r in para.runs[2:]:
            r.text = ""
    elif para.runs:
        label = para.runs[0].text.split(":")[0] + ": "
        para.runs[0].text = label + new_value
    else:
        para.add_run(new_value)


def set_result(cell, dat: bool) -> None:
    yes, no = ("☑ Đạt", "☐ Chưa") if dat else ("☐ Đạt", "☑ Chưa")
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(yes)
        cell.add_paragraph(no)
        return
    set_run_text(paras[0], yes)
    if len(paras) == 1:
        cell.add_paragraph(no)
    else:
        set_run_text(paras[1], no)
        for p in paras[2:]:
            set_run_text(p, "")


def set_note(cell, text: str) -> None:
    paras = cell.paragraphs
    if not paras:
        p = cell.add_paragraph(text)
        for r in p.runs:
            r.font.size = Pt(9)
        return
    set_run_text(paras[0], text)
    for r in paras[0].runs:
        r.font.size = Pt(9)
    for p in paras[1:]:
        set_run_text(p, "")


def fill_row(tbl, ri: int, dat: bool, note: str) -> None:
    set_result(tbl.rows[ri].cells[2], dat)
    set_note(tbl.rows[ri].cells[3], note)


def fill_before() -> None:
    src = HERE / "BeforeSubmitting_h.docx"
    shutil.copy2(src, BACKUP / src.name)
    d = Document(str(src))
    d.paragraphs[4].text = "Tên bài báo: " + TITLE
    d.paragraphs[5].text = ""
    fill_after_label(d.paragraphs[6], AUTHORS_VN + " (" + AUTHORS_EN + ")")
    fill_after_label(d.paragraphs[7], NCS)
    fill_after_label(d.paragraphs[8], VENUE)
    fill_after_label(d.paragraphs[9], DATE_VN + " (bản A29, pack OJS A30)")
    set_run_text(d.paragraphs[12], DATE_LINE)

    t = d.tables[0]
    # MỨC 1
    fill_row(
        t, 2, True,
        "Đã đối chiếu [1]–[26] với Crossref/doi.org và trang publisher (audit 31/08/2026: "
        "IJIET_FINAL_REVISION + bản lưu REFERENCE_LIVE_AUDIT). Mọi mục là công trình thật; "
        "không bịa DOI. File PDF gốc của từng bài NCS lưu ngoài repo (audit không chứa full PDF).",
    )
    fill_row(
        t, 3, True,
        "Không claim “first/SOTA/outperforms all”. SimpleKT [4] chỉ là related work, không phải T-KT. "
        "UKT [21], Mitton [22], AIED 2025 [17] đã đưa vào literature. Claim mạnh đều gắn Table 4–6 / S4.",
    )
    fill_row(
        t, 4, True,
        "Related work và Methods đã viết lại theo protocol TSCDA (không copy introduction từ bài khác). "
        "Chưa có báo cáo Turnitin — xem Mức 4 mục 2 (đánh Chưa).",
    )
    fill_row(
        t, 5, True,
        "IJIET ethics §6.3: Generative AI Statement + cover letter. Công cụ: ChatGPT GPT-5.6, "
        "Claude Sonnet 5, Google Antigravity 2.11.0, Cursor Grok 4.6 (public ID ngày 01/09/2026). "
        "AI không đồng tác giả; không bịa số liệu.",
    )
    # MỨC 2
    fill_row(
        t, 7, True,
        "Công thức chính: ECE (1), Brier/Murphy, FAR=P(y=0|p≥τ), Excess FAR. Không có kiến trúc mới "
        "cần chứng minh. Chiều tensor T-KT: Transformer 2 lớp, 4 heads, embed 64 (Table 2).",
    )
    fill_row(
        t, 8, True,
        "Text–công thức–bảng khớp: τ=0.7; FAR/Miss/Nadvance. CSV class simplekt = T-KT trong bài "
        "(không phải SimpleKT [4]). Occupancy R/L/I theo N test.",
    )
    fill_row(
        t, 9, True,
        "Không leakage tần suất KC: f_train chỉ trên train. Split learner-based; 5 seed "
        "(42, 2024–2027), 4 partition (fold_2=fold_3). Metric: AUC/ECE/FAR (không RecSys@K). "
        "Báo cáo mean±sd trên 4 partition. Code trong code_for_review_anonymous.zip.",
    )
    fill_row(
        t, 10, True,
        "Không dùng “significant” không có kiểm định. Gate dùng CI KC-cluster B=2000: "
        "ΔFAR T-KT [0.006, 0.138]. Bảng có ±sd. Không gọi chênh lệch nhỏ hơn nhiễu là SOTA.",
    )
    fill_row(
        t, 11, True,
        "Thuật ngữ neo literature: ECE [12], Brier [13], FAR (định nghĩa trong bài), occupancy R/L/I "
        "(descriptive). TSCDA = Train-only Sparse-Concept Decision Audit, định nghĩa ở abstract/III.D. "
        "T-KT thống nhất toàn bài.",
    )
    fill_row(
        t, 12, True,
        "Không overclaim: 3 dataset; ASSISTments Limited N=415 / Nadvance=235; Junyi sparse empty; "
        "XES ECE phẳng. Có mục V.D Limitations. Không classroom RCT. Không “first”.",
    )
    fill_row(
        t, 13, True,
        "Không claim SOTA. Baseline chính IRT/DKT/T-KT; GKT+CL4KT exploratory 1 fold. "
        "Literature 2025–2026: [17][21][22]. Không thêm SimpleKT official / temperature trừ khi editor yêu cầu.",
    )
    fill_row(
        t, 14, True,
        "Table 1 in từ preprocessing: ASSISTments 27,806 / 265 / 2.66M / 534,150; Junyi 71,014 / "
        "1,326 / 16.2M; XES 18,066 / 865 / 6.41M (padding excluded, khớp [20] 865 KC / 7,652 item). "
        "Nguồn: ASSISTments 2012–13 [18], Junyi Kaggle [19], XES3G5M [20].",
    )
    # MỨC 3
    fill_row(
        t, 16, True,
        "II.A KT models; II.B GKT/CL4KT; II.C bốn nghĩa “sparse”; II.D calibration/decision. "
        "Gap: occupancy + ECE + locked-τ FAR cùng lúc — không phải bài kiến trúc mới.",
    )
    fill_row(
        t, 17, True,
        "Gap hẹp: population AUC che miscalibration trên KC thưa + quyết định τ. TSCDA từ chối claim "
        "khi bucket rỗng (Junyi ucid). Không bán thành luật phổ quát.",
    )
    fill_row(
        t, 18, True,
        "Rà 01/09/2026 trên PDF named+blind: ECE 0.1136/0.2280; FAR 0.196/0.268; dual N 415/444/413 "
        "không lẫn; XES 0.1176/0.1129/0.1254; title A27; 8 trang. JEDM = 0 hit trong article.",
    )
    fill_row(
        t, 19, True,
        "Đóng góp (i)–(iii) và phân tích Table 7–8 / S4 gắn log thật (ASSISTments seed-42 T-KT). "
        "AI chỉ language/format; số liệu từ analysis/ + a2b.",
    )
    fill_row(
        t, 20, True,
        "Không dùng lời khen của AI làm bằng chứng. Phiếu này gửi GS (PGS.TS. Nguyễn Văn Hậu) "
        "kèm gói OJS_UPLOAD trước khi nộp.",
    )
    fill_row(
        t, 21, True,
        "Khóa số từ artifact + PDF; DOI từ Crossref/publisher; không nhờ AI tự xác minh DOI/số liệu. "
        "Audit: SCIENTIFIC_LOCKS.md, FORMAT_A29, REFERENCE_LIVE_AUDIT.",
    )
    # MỨC 4
    fill_row(
        t, 23, False,
        "CHƯA CHẠY isgen.ai / GPTZero. Cần NCS quét PDF named, ghi % vào đây và lưu báo cáo. "
        "Không dùng công cụ “né máy dò”. Các đoạn bị đánh dấu cao phải đọc lại thủ công.",
    )
    fill_row(
        t, 24, False,
        "CHƯA CÓ báo cáo Turnitin/iThenticate của Trường. Similarity % để trống cho đến khi NCS chạy "
        "và lưu PDF báo cáo cùng hồ sơ. Related work đã viết lại (Mức 1 mục 3) nhưng chưa thay được tool report.",
    )
    fill_row(
        t, 25, True,
        "Crossref/doi.org + landing pages [1]–[26] (31/08/2026). [2][4][5][6][11][20][22] không có "
        "journal DOI Crossref — đúng, không bịa. [15] giữ arXiv:2606.14123. [19] năm 2019 vs Kaggle 2020 (nit).",
    )
    fill_row(
        t, 26, True,
        "Soát tiếng Anh bằng Cursor Grok 4.6 trên Word IJIET (không Grammarly). Thuật ngữ ECE/FAR/T-KT/TSCDA "
        "không bị công cụ đổi. Bản 8 trang, title 20 pt.",
    )
    fill_row(
        t, 27, True,
        "Số liệu bài khớp artifact đã khóa (không retrain 3 dataset×5 seed ngày 01/09). "
        "XES dùng series a2b masked. Zip review 95 file, không _archive, không named Word. "
        "README + requirements.txt trong repo.",
    )
    fill_row(
        t, 28, True,
        "Review: code_for_review_anonymous.zip. Sau review: github.com/trinhnkt/Sparse-Concept-and-Calibration. "
        "README đã viết lại theo IJIET (T-KT ≠ SimpleKT; không khung JEDM). Blind PDF không chứa URL GitHub.",
    )
    fill_row(
        t, 29, True,
        "Template IJIET_template.doc; A4 2-cột; 8 trang named + 8 trang blind; IEEE numbered refs. "
        "Gói nộp: IJIET_FINAL_REVISION/output/OJS_UPLOAD/ (full.doc/.pdf, blind.pdf, supplementary.pdf, "
        "zip, cover letter). Không upload _archive/.",
    )

    # signature hint
    sig = d.tables[1].rows[0].cells[0]
    extra = sig.paragraphs[-1] if sig.paragraphs else sig.add_paragraph()
    if "Nguyễn Khánh Trình" not in sig.text:
        p = sig.add_paragraph("Nguyễn Khánh Trình")
        p.runs[0].font.size = Pt(11) if p.runs else None

    d.save(str(src))
    print("filled", src)


def fill_congbo() -> None:
    src = HERE / "CongBoKhoaHoc_AInIS.docx"
    shutil.copy2(src, BACKUP / src.name)
    d = Document(str(src))
    d.paragraphs[4].text = "Tên bài báo: " + TITLE
    d.paragraphs[5].text = ""
    fill_after_label(d.paragraphs[6], AUTHORS_VN + " / " + AUTHORS_EN)
    # P7 has two fields
    runs = d.paragraphs[7].runs
    if len(runs) >= 4:
        runs[1].text = NCS
        runs[3].text = "NCS, first author (K.-T.N.)"
        for r in runs[4:]:
            r.text = ""
    else:
        fill_after_label(d.paragraphs[7], NCS + "     Vai trò: NCS, first author")
    fill_after_label(d.paragraphs[8], VENUE)
    # backup lines
    p9 = d.paragraphs[9]
    if len(p9.runs) >= 7:
        p9.runs[6].text = BACKUP1
        for r in p9.runs[7:]:
            r.text = ""
    else:
        set_run_text(p9, "Phương án dự phòng 1: " + BACKUP1)
    p10 = d.paragraphs[10]
    if len(p10.runs) >= 2:
        p10.runs[1].text = BACKUP2
        for r in p10.runs[2:]:
            r.text = ""
    else:
        set_run_text(p10, "2: " + BACKUP2)
    fill_after_label(d.paragraphs[11], DATE_VN)
    set_run_text(d.paragraphs[16], DATE_LINE)

    t = d.tables[0]
    fill_row(
        t, 2, True,
        "Thứ tự như bản named: K.-T.N., T.D.M., D.N.T., C.T.N., V.-H.N. (corresponding). "
        "CRediT trong bài. Cần GS và đồng tác giả xác nhận bản OJS_UPLOAD này trước khi upload.",
    )
    fill_row(
        t, 3, False,
        "Mục tiêu đã chốt: IJIET (scope EdTech / evaluation). Hai dự phòng phía trên là ĐỀ XUẤT — "
        "chưa thống nhất GS nên đánh Chưa. JEDM đã withdraw, không đặt lại làm dự phòng trừ khi GS quyết định.",
    )
    fill_row(
        t, 4, True,
        "Scopus Source ID 21100921050, ISSN 2010-3689; coverage Scopus 2019–2026 (không discontinued "
        "theo sourceid). Không lấy indexing chỉ từ www.ijiet.org. Clarivate MJL/WoS Core: IJIET không "
        "phải SCIE/SSCI — indexing chính là Scopus. Nhờ GS đối chiếu điều kiện công nhận của Trường.",
    )
    fill_row(
        t, 6, True,
        "Một manuscript / một venue: IJIET là bản đang active. Cover letter ghi JEDM đã withdraw. "
        "Article không nêu tên JEDM. NCS/GS giữ email xác nhận withdraw trước khi upload OJS.",
    )
    fill_row(
        t, 7, True,
        "Không salami: bản IJIET 8 trang là diagnostic instrument (TSCDA), không nộp song song. "
        "Số liệu bảng = log khóa (ECE 0.1136/0.2280; FAR 0.196/0.268; XES masked). "
        "Turnitin chưa chạy — xem phiếu AI Mức 4 mục 2.",
    )
    fill_row(
        t, 8, True,
        "Audit DOI 31/08/2026: mọi [1]–[26] tồn tại; citation khớp claim (T-KT ≠ SimpleKT [4]). "
        "Bibliography không do AI bịa. Chi tiết: REFERENCE_LIVE_AUDIT.md.",
    )
    fill_row(
        t, 9, True,
        "AI không phải tác giả. Khai báo §6.3 trong bài + cover letter (công cụ và phiên bản public "
        "01/09/2026). Tác giả chịu trách nhiệm số liệu và lập luận.",
    )
    fill_row(
        t, 11, False,
        "Đây là bước gửi GS bản submission-ready (OJS_UPLOAD). Chưa đủ 3 ngày + chữ ký đồng tác giả "
        "nên đánh Chưa cho đến khi GS xác nhận. Không đăng preprint khi chưa hỏi GS. Không mã đề tài "
        "trong bài (acknowledgment: institutions).",
    )
    fill_row(
        t, 12, True,
        "Blind PDF 8 trang: Anonymous Authors; metadata author rỗng; 0 hit Hung Yen / GitHub / email. "
        "Gói tái lập: zip anonymous + analysis locks + a2b scripts. Named Word chỉ cho editor. "
        "Chưa có submission ID (chưa upload).",
    )
    fill_row(
        t, 14, True,
        "Chưa gửi OJS nên chưa phát sinh withdraw/revision. Cam kết: không withdraw/chuyển venue khi "
        "chưa trao đổi GS; revision báo trong 24 giờ kèm decision letter; response Comment→Response→"
        "Change→Location; lưu reviewer comment vào sổ nhóm.",
    )

    sig = d.tables[1].rows[0].cells[0]
    if "Nguyễn Khánh Trình" not in sig.text:
        sig.add_paragraph("Nguyễn Khánh Trình")

    d.save(str(src))
    print("filled", src)


if __name__ == "__main__":
    fill_before()
    fill_congbo()

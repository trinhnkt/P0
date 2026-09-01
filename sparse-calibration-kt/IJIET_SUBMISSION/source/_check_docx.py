from docx import Document

p = r"C:\TRINH\Sparse-Concept and Calibration\sparse-calibration-kt\ijiet\Reproducible_Sparse_Concept_and_Calibration_Diagnostics_for_Knowledge_Tracing.docx"
d = Document(p)
out = open(
    r"C:\TRINH\Sparse-Concept and Calibration\sparse-calibration-kt\IJIET_SUBMISSION\audit\docx_char_check.txt",
    "w",
    encoding="utf-8",
)
for para in d.paragraphs:
    t = para.text
    if "0.1136" in t or "mean" in t.lower() and "sd" in t.lower() and "Table 2" in t:
        out.write(repr(t) + "\n")
        out.write(" ".join(f"U+{ord(c):04X}" for c in t[t.find('0.1136')-5:t.find('0.1136')+20] if "0.1136" in t) + "\n")
    if "Table 2." in t:
        out.write("T2 " + repr(t) + "\n")
    if "Table 4." in t:
        out.write("T4 " + repr(t) + "\n")
# refs
nref = 0
for para in d.paragraphs:
    if para.style.name == "References" and para.text.strip():
        nref += 1
        if nref <= 3 or nref >= 18:
            out.write(f"REF{nref} {para.text[:100]}\n")
out.write(f"nref={nref}\n")
# table cell sample
if d.tables:
    out.write("T2r2c3=" + repr(d.tables[1].rows[1].cells[2].text) + "\n")
    out.write("T2r2c4=" + repr(d.tables[1].rows[1].cells[3].text) + "\n")
out.close()
print("wrote char check")

from docx import Document
from pathlib import Path

p = Path(r"C:\TRINH\Sparse-Concept and Calibration\sparse-calibration-kt\IJIET_SUBMISSION\source\main_ijiet_step03.docx")
d = Document(str(p))
out = Path(r"C:\TRINH\Sparse-Concept and Calibration\sparse-calibration-kt\IJIET_SUBMISSION\audit\step03_front.txt")
lines = []
for i, para in enumerate(d.paragraphs[:15], 1):
    t = para.text
    if t.strip() or i <= 8:
        lines.append(f"{i} [{para.style.name}] {t[:200]!r}")
text_all = "\n".join(para.text for para in d.paragraphs[:20])
lines.append("HAS_INDEX_TERMS=" + str("Index Terms" in text_all))
lines.append("HAS_KEYWORDS_EM=" + str("Keywords—" in text_all or "Keywords\u2014" in text_all))
lines.append("HAS_GATES_TITLE=" + str("Remediation Gates" in text_all))
lines.append("HAS_PREFERRED=" + str("Threshold-Based Educational Decisions" in text_all))
lines.append("HAS_JEDM_TITLE=" + str("Reproducible Sparse-Concept" in d.paragraphs[0].text))
# abstract still present
for para in d.paragraphs:
    if para.text.startswith("Abstract"):
        lines.append("ABSTRACT_START=" + repr(para.text[:80]))
        break
out.write_text("\n".join(lines) + "\n", encoding="utf-8")
print(out.read_text(encoding="utf-8"))

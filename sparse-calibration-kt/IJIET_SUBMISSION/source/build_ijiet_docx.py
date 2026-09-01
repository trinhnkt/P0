#!/usr/bin/env python3
"""Build an IJIET-styled Word manuscript from validated IJIET draft content.

Formatting authority: current published IJIET articles (Vol. 16, 2026) and
author guide https://www.ijiet.org/list-14-1.html, because
https://www.ijiet.org/files/IJIET_template.doc returned HTTP 500.

Scientific numbers are copied from ijiet/main_ijiet.tex and
analysis/four_partition/. No results are recomputed.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "ijiet"
OUT_NAME = "Reproducible_Sparse_Concept_and_Calibration_Diagnostics_for_Knowledge_Tracing.docx"
OUT_PATH = OUT_DIR / OUT_NAME
COPY_PATH = ROOT / "IJIET_SUBMISSION" / "output" / OUT_NAME
FIG_CANDIDATES = [
    ROOT / "paper" / "figures" / "figure2_bucket_distribution.png",
    ROOT / "paper" / "figures" / "figure2_bucket_distribution.pdf",
    ROOT / "REV_REVIEWER_CALIBRATION_v1" / "figures" / "figure2_bucket_distribution.png",
]


def set_run_font(run, name="Times New Roman", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(p, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0,
                         line=12, first_line=None, space_after_exact=None):
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after if space_after_exact is None else space_after_exact)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line)
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    else:
        pf.first_line_indent = Cm(0)


def add_text(p, text, **font):
    run = p.add_run(text)
    set_run_font(run, **font)
    return run


def add_para(doc, text, *, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=0, line=12, first_line=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=align, before=before, after=after, line=line, first_line=first_line)
    if text:
        add_text(p, text, size=size, bold=bold, italic=italic)
    return p


def add_mixed_para(doc, parts, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0,
                   line=12, first_line=None):
    """parts: list of (text, kwargs_for_font)."""
    p = doc.add_paragraph()
    set_paragraph_format(p, align=align, before=before, after=after, line=line, first_line=first_line)
    for text, kwargs in parts:
        add_text(p, text, **kwargs)
    return p


def add_section_heading(doc, text):
    return add_para(
        doc, text,
        size=10, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=12, after=4, line=12,
    )


def add_subsection_heading(doc, text):
    return add_para(
        doc, text,
        size=10, italic=True, bold=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        before=8, after=3, line=12,
    )


def add_body(doc, text):
    return add_para(doc, text, size=10, first_line=0.42, before=0, after=0, line=12)


def set_two_columns(section, ncols=2, space_twips=360):
    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn("w:cols"):
            sectPr.remove(child)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(ncols))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


def set_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    set_paragraph_format(hp, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=11)
    add_text(
        hp,
        "International Journal of Information and Education Technology",
        size=9, italic=True,
    )

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    set_paragraph_format(fp, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=11)
    add_text(fp, "Copyright © 2026 by the authors. ", size=8, italic=True)
    add_text(fp, "CC BY 4.0", size=8, italic=True)

    # PAGE field
    p = section.footer.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=11)
    run = p.add_run()
    set_run_font(run, size=9)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def shade_cell(cell, hex_color="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def fill_table(table, rows, header=True, font_size=8):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=1, line=10)
            add_text(p, str(val), size=font_size, bold=(header and i == 0))
            set_cell_border(cell)
            if header and i == 0:
                shade_cell(cell, "F2F2F2")


def add_caption(doc, text, *, before=6, after=4):
    return add_para(
        doc, text, size=9, bold=False, italic=False,
        align=WD_ALIGN_PARAGRAPH.CENTER, before=before, after=after, line=11,
    )


def prevent_row_break(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)


def find_figure():
    for p in FIG_CANDIDATES:
        if p.exists() and p.suffix.lower() == ".png":
            return p
    return None


def build():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    section.different_first_page_header_footer = False
    set_header_footer(section)

    # Normal style
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # ----- Front matter (full width) -----
    add_para(
        doc,
        "Reproducible Sparse-Concept and Calibration Diagnostics for Knowledge Tracing",
        size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=10, line=26,
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2, line=14)
    add_text(p, "Khanh-Trinh Nguyen", size=11)
    add_text(p, "1", size=9)
    add_text(p, ", Tuan Dao Minh", size=11)
    add_text(p, "1", size=9)
    add_text(p, ", Duong Nguyen Tien", size=11)
    add_text(p, "1", size=9)
    add_text(p, ", Chi Thanh Nguyen", size=11)
    add_text(p, "2", size=9)
    add_text(p, ", and Van-Hau Nguyen", size=11)
    add_text(p, "1, *", size=9)

    add_para(
        doc,
        "1. Hung Yen University of Technology and Education, Hung Yen, Vietnam",
        size=9, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=0, line=11,
    )
    add_para(
        doc,
        "2. Academy of Military Science and Technology, Ha Noi, Vietnam",
        size=9, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2, line=11,
    )
    add_para(
        doc,
        "Email: trinhnk@utehy.edu.vn (K.-T.N.); tuanymc@utehy.edu.vn (T.D.M.); "
        "duongnt@utehy.edu.vn (D.N.T.); thanhnc@ioit.ai.vn (C.T.N.); "
        "haunv@utehy.edu.vn (V.-H.N.)",
        size=8, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=11,
    )
    add_para(
        doc, "*Corresponding author",
        size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6, line=11,
    )
    add_para(
        doc,
        "Manuscript received XX Month 2026; revised XX Month 2026; "
        "accepted XX Month 2026; published XX Month 2026",
        size=8, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8, line=10,
    )

    # Switch to two columns
    new_section = doc.add_section()
    new_section.start_type = WD_SECTION.CONTINUOUS
    new_section.page_width = Inches(8.5)
    new_section.page_height = Inches(11)
    new_section.top_margin = Inches(0.75)
    new_section.bottom_margin = Inches(1.0)
    new_section.left_margin = Inches(0.625)
    new_section.right_margin = Inches(0.625)
    new_section.header_distance = Inches(0.4)
    new_section.footer_distance = Inches(0.4)
    set_two_columns(new_section, 2, space_twips=360)
    set_header_footer(new_section)

    # Abstract
    add_mixed_para(
        doc,
        [
            ("Abstract", {"size": 9, "bold": True, "italic": False}),
            ("—", {"size": 9, "bold": True}),
            (
                "Knowledge Tracing (KT) systems that drive practice, remediation, or "
                "advancement typically consume a predicted probability, not an "
                "area-under-the-curve (AUC) score. Aggregate ranking metrics can look "
                "acceptable while those probabilities are poorly calibrated on rarely "
                "practiced skills. This paper reports a diagnostic evaluation—not a new "
                "KT architecture—of how calibration behaves across train-only "
                "knowledge-component (KC) frequency strata on three public logs "
                "(ASSISTments 2012, Junyi Academy, and XES3G5M) with IRT, DKT, and "
                "SimpleKT. Sparse training evidence does not universally reduce "
                "discrimination. It can, however, expose a dataset-dependent calibration "
                "vulnerability. On ASSISTments 2012, SimpleKT expected calibration error "
                "(ECE) rises from 0.114 on dense KCs to 0.228 on sparse KCs (Limited "
                "occupancy, N=415); Junyi has no learner-based sparse stratum, and "
                "XES3G5M SimpleKT ECE is essentially flat. Treating the same "
                "probabilities as a locked global mastery gate at τ=0.7 raises SimpleKT "
                "false mastery among advances from 0.196 (dense) to 0.268 (sparse) on "
                "one ASSISTments fold; the sparse–dense gap stays positive on all five "
                "training seeds (mean 0.047; four unique student partitions). A "
                "train-only graph KT instantiation shrinks that seed-42 gap. The gate is "
                "a simulation, not a classroom trial. The practical implication for "
                "educational-technology designers is to validate thresholded decisions "
                "per frequency stratum rather than on population AUC or ECE alone.",
                {"size": 9},
            ),
        ],
        before=0, after=6, line=11,
    )
    add_mixed_para(
        doc,
        [
            ("Keywords", {"size": 9, "bold": True}),
            ("—", {"size": 9, "bold": True}),
            (
                "knowledge tracing, learning analytics, calibration, sparse concepts, "
                "mastery threshold, educational decision support",
                {"size": 9, "italic": True},
            ),
        ],
        before=0, after=8, line=11,
    )

    # I. INTRODUCTION
    add_section_heading(doc, "I. INTRODUCTION")
    add_body(
        doc,
        "Adaptive practice engines and intelligent tutoring systems increasingly use "
        "Knowledge Tracing (KT) to estimate whether a learner has mastered a skill and "
        "to decide whether to skip, remediate, or advance [1]–[3]. In deployment, that "
        "decision is rarely the AUC of a leaderboard. It is a probability p compared "
        "with a threshold τ: if p is high enough, the system treats the skill as "
        "mastered and withholds extra help.",
    )
    add_body(
        doc,
        "KT papers typically report population AUC and accuracy [4], [5]. Those metrics "
        "are useful for ranking models, but they can hide reliability problems on "
        "low-frequency knowledge components (KCs)—concepts with little training-fold "
        "evidence—because most test events sit on dense, frequently practiced skills. A "
        "model can rank well overall and still be overconfident on the sparse tail. "
        "When p is used as a mastery or remediation gate, that overconfidence is not a "
        "cosmetic reporting issue: it is a false-mastery error.",
    )
    add_body(
        doc,
        "This paper treats that mismatch as an evaluation problem for educational "
        "technology, not as a reason to propose another KT backbone. We ask three "
        "questions that a system designer can act on: (1) Does lower KC training "
        "frequency systematically worsen ranking (AUC) across datasets? (2) Does it "
        "worsen probability reliability (calibration), and is that pattern "
        "dataset-dependent? (3) If a global threshold is applied to p, does the sparse "
        "stratum produce a dirtier advance decision than the dense stratum?",
    )
    add_body(
        doc,
        "The empirical story is bounded. Sparse training evidence does not universally "
        "degrade discrimination. On XES3G5M, sparse AUC is higher than dense AUC for "
        "DKT and SimpleKT. Calibration, however, can be more sensitive than aggregate "
        "ranking on some logs. On ASSISTments 2012, SimpleKT ECE increases from dense "
        "to sparse concepts, and a locked gate at τ=0.7 produces a higher false-mastery "
        "rate among sparse advances. The same ECE gradient is absent on Junyi (empty "
        "sparse bucket) and essentially absent for SimpleKT on XES3G5M. We therefore "
        "do not claim that “sparse KCs always fail.” We claim that occupancy-aware, "
        "stratum-wise calibration—and, when probabilities are consumed as gates, a "
        "simulated decision error—should be checked before a KT model is wired into "
        "remediation logic.",
    )
    add_body(
        doc,
        "Contributions, stated conservatively: (i) a train-only KC-frequency protocol "
        "with an explicit strict cold-start group, so sparse diagnostics cannot leak "
        "test-fold counts into the definition of “sparse”; (ii) per-stratum calibration "
        "(ECE, Brier decomposition) on three public datasets, with sample-size flags "
        "(Reliable / Limited / Insufficient); (iii) a locked-threshold simulation of "
        "false mastery and miss rates, plus a five-seed check of the sparse–dense "
        "false-mastery gap on ASSISTments 2012. We do not claim a new audit theory, a "
        "new calibrator, or a classroom RCT. Graph KT (GKT) and a CL4KT-style adapter "
        "appear only as a single-fold diagnostic instantiation on ASSISTments, not as "
        "a model-contribution bake-off.",
    )

    # II. RELATED WORK
    add_section_heading(doc, "II. RELATED WORK")
    add_body(
        doc,
        "Classical KT uses Bayesian knowledge tracing or item-response models [1], [6]. "
        "Deep sequential models (DKT, attentive and transformer KT) improved aggregate "
        "prediction [2], [4], [7]. Graph and contrastive variants often motivate "
        "architecture by sparse or related skills [8], [9]. Those papers still report "
        "mainly population metrics.",
    )
    add_body(
        doc,
        "Educational measurement has long warned that ranking is not reliability [10]. "
        "Calibration error and reliability diagrams quantify whether predicted "
        "probabilities match empirical correctness [11]–[14]. Recent KT work has begun "
        "to study calibration and sparse-concept evaluation [15]–[17], but deployed "
        "tutoring systems still typically consume a single global p. This study sits in "
        "that gap: we keep standard models, stratify by training-only frequency, and "
        "ask whether a population gate would treat sparse skills more leniently—or more "
        "dangerously—than dense ones.",
    )

    # III. METHOD
    add_section_heading(doc, "III. METHOD")
    add_subsection_heading(doc, "A. Datasets, Splits, and Models")
    add_body(
        doc,
        "We use three de-identified public logs after a common schema: ASSISTments "
        "2012 [18], [5], Junyi Academy [19], and XES3G5M [20]. Table 1 summarizes "
        "processed cohort size. Learner-based splits (unseen students) are the primary "
        "setting. Temporal splits are a complementary stress test and are not the "
        "source of the gate numbers below.",
    )

    add_caption(doc, "Table 1. Processed cohort statistics (learner-based split)", before=8, after=2)
    t1 = doc.add_table(rows=4, cols=5)
    fill_table(
        t1,
        [
            ["Dataset", "Learners", "KCs", "Interactions", "Test events"],
            ["ASSISTments 2012", "27,806", "265", "2.66M", "534,150"],
            ["Junyi Academy", "71,014", "1,326", "16.2M", "3,269,022"],
            ["XES3G5M", "18,066", "866", "7.95M", "1,589,145"],
        ],
        font_size=8,
    )
    prevent_row_break(t1)

    add_body(
        doc,
        "Baselines: one-parameter IRT [6] as a classical reference, DKT [2], and "
        "SimpleKT [4]. On ASSISTments fold 0 only we also score a train-only GKT "
        "graph [8] and a CL4KT protocol adapter [9] (contrastive views on training "
        "sequences; not an official CL4KT checkpoint). IRT under learner-based splits "
        "has no ability parameter for unseen students, so its AUC is 0.50 by "
        "construction; we report it as a base-rate reference, not as a ranking "
        "competitor.",
    )
    add_body(
        doc,
        "Deep models are trained at seeds 42, 2024, 2025, 2026, and 2027. The processed "
        "partitions for seeds 2025 and 2026 are identical (fold_2 = fold_3) on all three "
        "datasets. Tables that report mean±sd therefore summarize four unique student "
        "partitions: the two initializations on the duplicated split are averaged first. "
        "The gate-robustness table still lists all five training seeds without dropping "
        "one after the fact.",
    )

    add_subsection_heading(doc, "B. Train-Only Frequency Strata")
    add_body(
        doc,
        "For each fold, KC frequency f_train is counted on the training file only. "
        "Buckets: strict cold-start (f=0), very sparse (0<f<20), sparse (20≤f<100), "
        "medium (100≤f<500), dense (f≥500). Dense KCs dominate event volume while a "
        "non-empty sparse-like tail exists on ASSISTments and XES3G5M. Occupancy flags "
        "follow test-event count N: Reliable (N≥1000), Limited (100≤N<1000), "
        "Insufficient (N<100). Success claims require Limited or Reliable occupancy. "
        "Very-sparse ASSISTments cells are Insufficient and descriptive only.",
    )

    fig = find_figure()
    if fig is not None:
        p = doc.add_paragraph()
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2, line=12)
        run = p.add_run()
        run.add_picture(str(fig), width=Inches(3.35))
        add_caption(
            doc,
            "Fig. 1. Train-only KC-frequency strata. Dense concepts dominate "
            "interactions; sparse and cold-start KCs remain the diagnostic tail.",
            before=2, after=6,
        )
    else:
        add_para(
            doc,
            "Fig. 1 (source file paper/figures/figure2_bucket_distribution.pdf) is "
            "embedded in the compiled PDF; it was not copied here because the raster "
            "asset was not available to this builder.",
            size=8, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6, line=10,
        )

    add_subsection_heading(doc, "C. Calibration")
    add_body(
        doc,
        "Let y∈{0,1} be correctness and p∈[0,1] the predicted probability. With M=15 "
        "equal-width bins,",
    )
    add_para(
        doc,
        "ECE = Σ_m (n_m / N) |acc_m − conf_m|.",
        size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=4, line=14,
    )
    add_body(
        doc,
        "We also report the Brier score and its reliability / resolution / uncertainty "
        "decomposition [14]. Lower ECE is better calibration; it is not a substitute "
        "for AUC.",
    )

    add_subsection_heading(doc, "D. Simulated Mastery Gate")
    add_body(
        doc,
        "For a locked global threshold τ, the system advances (skips remediation) if "
        "p≥τ and triggers remediation otherwise. Sparse events never receive their own "
        "tuned τ. The display threshold is τ=0.7; the grid {0.5, 0.6, 0.7, 0.8} is "
        "recorded in the artifact and is not a search over sparse error.",
    )
    add_body(
        doc,
        "Let A={p≥τ}. False mastery (FM) is P(y=0 | A): among advances, how often was "
        "the answer wrong. Miss is P(A | y=0): among wrong answers, how often the "
        "system still skipped help. If the model were calibrated on the advance set, "
        "FM would match E[1−p | A]. We report ΔFM = FM_sparse − FM_dense. Positive ΔFM "
        "means sparse advances are dirtier than dense advances. This is a simulated "
        "decision error, not an instructional RCT.",
    )

    # IV. RESULTS
    add_section_heading(doc, "IV. RESULTS")
    add_subsection_heading(doc, "A. Aggregate Ranking Is Not the Sparse Story")
    add_body(
        doc,
        "Table 2 shows learner-based AUC. DKT is slightly above SimpleKT on all three "
        "datasets. IRT AUC is 0.50 for the reason given above. These numbers would look "
        "like a routine KT bake-off. They do not tell a designer whether p is "
        "trustworthy on the tail.",
    )

    add_caption(
        doc,
        "Table 2. Overall learner-based AUC (mean±sd over four unique partitions)",
        before=8, after=2,
    )
    t2 = doc.add_table(rows=10, cols=4)
    fill_table(
        t2,
        [
            ["Dataset", "Model", "AUC", "ACC"],
            ["ASSISTments 2012", "IRT", "0.5000", "0.6973±0.0004"],
            ["ASSISTments 2012", "DKT", "0.6979±0.0014", "0.7182±0.0014"],
            ["ASSISTments 2012", "SimpleKT", "0.6837±0.0025", "0.6996±0.0032"],
            ["Junyi Academy", "IRT", "0.5000", "0.7053±0.0018"],
            ["Junyi Academy", "DKT", "0.7320±0.0009", "0.7343±0.0015"],
            ["Junyi Academy", "SimpleKT", "0.7231±0.0030", "0.7274±0.0015"],
            ["XES3G5M", "IRT", "0.5000", "0.7961±0.0031"],
            ["XES3G5M", "DKT", "0.8171±0.0022", "0.8327±0.0032"],
            ["XES3G5M", "SimpleKT", "0.7557±0.0013", "0.8067±0.0037"],
        ],
        font_size=8,
    )
    prevent_row_break(t2)

    add_body(
        doc,
        "On XES3G5M, sparse AUC is higher than dense AUC (DKT 0.857 vs. 0.817; "
        "SimpleKT 0.847 vs. 0.755; Reliable occupancy). Lower frequency is therefore "
        "not a universal ranking failure. Junyi’s learner-based sparse bucket is empty "
        "under the pre-registered cuts—a protocol outcome, not a missing cell to impute.",
    )

    add_subsection_heading(doc, "B. Calibration Can Move When Ranking Does Not")
    add_body(
        doc,
        "Table 3 is the calibration punchline. On ASSISTments 2012, SimpleKT ECE "
        "increases from 0.1136±0.0066 (dense, Reliable, N=523,971) to 0.1541 (medium) "
        "to 0.2280±0.0197 (sparse, Limited, N=415). DKT shows the same direction "
        "(dense ECE 0.060 vs. sparse 0.233). IRT dense ECE is tiny (0.003) but "
        "resolution is zero: a constant-like predictor can look “calibrated” while "
        "ranking nothing.",
    )
    add_body(
        doc,
        "On Junyi, only dense and medium strata exist; SimpleKT ECE rises modestly from "
        "0.079 to 0.107. On XES3G5M, SimpleKT ECE is essentially flat "
        "(0.114→0.111→0.125 dense/medium/sparse) despite Reliable sparse occupancy "
        "(N=2,010). ECE-flat is not a license to skip occupancy reporting, and—as the "
        "gate results show—it is not the same as a flat miss rate.",
    )

    add_caption(
        doc,
        "Table 3. SimpleKT event-level ECE by train-only frequency stratum "
        "(four unique partitions). Flags: R Reliable, L Limited. Junyi sparse is empty.",
        before=8, after=2,
    )
    t3 = doc.add_table(rows=9, cols=5)
    fill_table(
        t3,
        [
            ["Dataset", "Stratum", "N", "Flag", "ECE"],
            ["ASSISTments 2012", "dense", "523,971", "R", "0.1136±0.0066"],
            ["ASSISTments 2012", "medium", "5,963", "R", "0.1541±0.0051"],
            ["ASSISTments 2012", "sparse", "415", "L", "0.2280±0.0197"],
            ["Junyi Academy", "dense", "3,232,614", "R", "0.0792±0.0051"],
            ["Junyi Academy", "medium", "3,836", "R", "0.1073±0.0156"],
            ["XES3G5M", "dense", "1,268,696", "R", "0.1145±0.0011"],
            ["XES3G5M", "medium", "12,980", "R", "0.1114±0.0076"],
            ["XES3G5M", "sparse", "2,010", "R", "0.1248±0.0085"],
        ],
        font_size=8,
    )
    prevent_row_break(t3)

    add_subsection_heading(doc, "C. A Global Gate Can Turn the ASSISTments Gradient into a Decision Error")
    add_body(
        doc,
        "Table 4 applies τ=0.7 on ASSISTments fold 0 (seed 42; sparse N=444, Limited). "
        "SimpleKT false mastery among advances is 0.196 on dense KCs and 0.268 on "
        "sparse KCs (ΔFM=+0.072). The calibrated expectation E[FM] on sparse advances "
        "is only 0.050, so the excess error is large. DKT shows a similar FM gap. "
        "Train-only GKT shrinks ΔFM to +0.015 (FM 0.205→0.220). The CL4KT adapter is "
        "intermediate (0.185→0.240). We do not read this as “use GKT in production”; "
        "we read it as: the ASSISTments calibration gradient is not a property of the "
        "dataset alone.",
    )

    add_caption(
        doc,
        "Table 4. Simulated gate at τ=0.7, ASSISTments 2012 fold 0. Advance if p≥τ. "
        "FM=P(y=0 | p≥τ); E[FM]=E[1−p | A]. Not a classroom trial.",
        before=8, after=2,
    )
    t4 = doc.add_table(rows=9, cols=6)
    fill_table(
        t4,
        [
            ["Model", "Stratum", "N", "FM", "E[FM]", "Miss"],
            ["SimpleKT", "dense", "528,018", "0.196", "0.113", "0.352"],
            ["SimpleKT", "sparse", "444", "0.268", "0.050", "0.320"],
            ["DKT", "dense", "528,018", "0.200", "0.147", "0.398"],
            ["DKT", "sparse", "444", "0.296", "0.057", "0.365"],
            ["GKT (train-only)", "dense", "528,018", "0.205", "0.163", "0.455"],
            ["GKT (train-only)", "sparse", "444", "0.220", "0.157", "0.234"],
            ["CL4KT (adapter)", "dense", "528,018", "0.185", "0.175", "0.359"],
            ["CL4KT (adapter)", "sparse", "444", "0.240", "0.116", "0.244"],
        ],
        font_size=7,
    )
    prevent_row_break(t4)

    add_body(
        doc,
        "Table 5 checks whether SimpleKT ΔFM is a one-fold accident. It is positive on "
        "all five training seeds (mean 0.047, sd 0.033). A KC-clustered bootstrap on "
        "seed 42 yields a 95% interval [0.006, 0.138]—excluding 0, but wide, as Limited "
        "occupancy requires. DKT ΔFM is positive on only three of five seeds and is "
        "not treated as a five-run finding. On XES3G5M, SimpleKT ΔFM is negative on all "
        "five seeds, but ΔMiss is positive on all five (mean +0.112): among actual "
        "incorrect answers, the system still advances more often on sparse than on "
        "dense KCs. A flat ECE therefore does not imply a flat miss rate.",
    )

    add_caption(
        doc,
        "Table 5. Gate robustness at τ=0.7 on ASSISTments 2012 (five training seeds; "
        "four unique partitions). GKT/CL4KT remain seed 42 only.",
        before=8, after=2,
    )
    t5 = doc.add_table(rows=3, cols=5)
    fill_table(
        t5,
        [
            ["Model", "Mean ΔFM", "SD", "Seeds >0", "Mean sparse N"],
            ["SimpleKT", "0.047", "0.033", "5/5", "413"],
            ["DKT", "0.033", "0.048", "3/5", "413"],
        ],
        font_size=8,
    )
    prevent_row_break(t5)

    add_body(
        doc,
        "Coincidence of digits: SimpleKT dense E[FM] at τ=0.7 on seed 42 is 0.113, "
        "close to the four-partition dense ECE 0.114. They are different quantities.",
    )

    # V. DISCUSSION
    add_section_heading(doc, "V. DISCUSSION")
    add_subsection_heading(doc, "A. What a Designer Should Take Away")
    add_body(
        doc,
        "If a learning platform uses KT probabilities only to sort students, Table 2 "
        "may suffice. If it uses p to skip practice or declare mastery, Table 2 is the "
        "wrong dashboard. The ASSISTments SimpleKT cell is the cautionary case: "
        "competitive AUC, Limited-but-estimable sparse occupancy (about 19% of KCs fall "
        "below the sparse threshold on fold 0), a monotonic ECE rise, and a gate that "
        "advances more dirty sparse attempts. The XES3G5M cell is the opposite caution: "
        "plenty of sparse mass and Reliable occupancy, yet SimpleKT ECE does not "
        "degrade—while miss rates still can. Junyi shows that the protocol can refuse a "
        "sparse claim when the bucket is empty.",
    )
    add_body(
        doc,
        "Three observable conditions track when sparse-calibration claims are "
        "informative rather than obligatory (Table 6): (1) a non-empty low-frequency "
        "tail under the split actually used; (2) enough sparse test events for at least "
        "a Limited-flag ECE; (3) frequency–difficulty coupling that is not inverted. "
        "ASSISTments meets all three. That is a diagnostic pre-condition, not a claim "
        "that training frequency causes miscalibration.",
    )

    add_caption(
        doc,
        "Table 6. When a sparse-calibration claim is estimable (findings, not causes). "
        "Sparse mass: share of KCs with f_train<100.",
        before=8, after=2,
    )
    t6 = doc.add_table(rows=5, cols=4)
    fill_table(
        t6,
        [
            ["Condition", "ASSISTments", "Junyi", "XES3G5M"],
            ["Sparse mass", "18.9% of KCs", "0%", "22.5% of KCs"],
            ["Sparse N", "415 (L)", "empty", "2,010 (R)"],
            ["Difficulty coupling", "ρ=−0.227 (sparse harder)", "ρ=−0.416", "ρ=+0.087 (weak, inverted)"],
            ["SimpleKT ECE", "0.114→0.228", "dense→medium only", "flat 0.114→0.125"],
        ],
        font_size=7,
    )
    prevent_row_break(t6)

    add_subsection_heading(doc, "B. What This Paper Does Not Show")
    add_body(
        doc,
        "The gate is not a classroom policy and not an A/B test of remediation. GKT and "
        "the CL4KT adapter are single-fold, ASSISTments-only instantiations. Temporal "
        "results are a single corrected cutoff (seed 42), not a multi-cutoff variance "
        "estimate. IRT is not a fair ranking baseline on unseen learners. We did not "
        "tune τ on sparse events, and we do not recommend doing so in deployment "
        "without a separate validation design.",
    )

    add_subsection_heading(doc, "C. Implications for Information and Education Technology")
    add_body(
        doc,
        "IJIET’s audience includes operators of educational platforms, not only KT "
        "researchers. For that audience, the operational recommendation is narrow: "
        "(i) log train-only KC frequency with every prediction; (ii) report ECE and "
        "occupancy on dense vs. sparse slices before enabling a global mastery "
        "threshold; (iii) if a threshold must be global, simulate FM and Miss on the "
        "sparse slice at the same τ used in the product; (iv) do not treat a population "
        "AUC win, or a graph/contrastive architecture, as automatically safer on the "
        "tail.",
    )

    # VI. CONCLUSION
    add_section_heading(doc, "VI. CONCLUSION")
    add_body(
        doc,
        "KT models that look adequate on AUC can still be poorly calibrated on rarely "
        "practiced skills, and a global mastery threshold can then produce a higher "
        "false-mastery rate on those skills. That pattern appears for SimpleKT on "
        "ASSISTments 2012 (ECE 0.114→0.228, Limited N=415; seed-42 FM 0.196→0.268; "
        "ΔFM>0 on 5/5 seeds). It does not appear as a universal law: Junyi has no "
        "learner-based sparse stratum, and XES3G5M SimpleKT ECE is flat. Sparse-concept "
        "diagnostics are therefore conditionally important. The simulation is a "
        "decision-error check for educational-technology systems, not a claim of a new "
        "KT model and not a classroom intervention.",
    )

    add_section_heading(doc, "CONFLICT OF INTEREST")
    add_body(doc, "The authors declare no conflict of interest.")

    add_section_heading(doc, "AUTHOR CONTRIBUTIONS")
    add_body(
        doc,
        "K.-T.N. led the diagnostic design, experiments, and manuscript. T.D.M. and "
        "D.N.T. contributed data processing and baseline runs. C.T.N. contributed "
        "methodological review. V.-H.N. supervised the study and revised the "
        "manuscript. All authors approved the final version.",
    )

    add_section_heading(doc, "ACKNOWLEDGMENT")
    add_body(
        doc,
        "Anonymized code and prediction exports are available for peer review; a public "
        "repository will be released upon acceptance. This study uses de-identified "
        "public learner logs (ASSISTments 2012, Junyi Academy, XES3G5M) and is not a "
        "classroom trial.",
    )

    add_section_heading(doc, "REFERENCES")
    refs = [
        '[1] A. T. Corbett and J. R. Anderson, “Knowledge Tracing: Modeling the acquisition of procedural knowledge,” User Modeling and User-Adapted Interaction, vol. 4, no. 4, pp. 253–278, 1994.',
        '[2] C. Piech, J. Bassen, J. Huang, S. Ganguli, M. Sahami, L. J. Guibas, and J. Sohl-Dickstein, “Deep Knowledge Tracing,” in Advances in Neural Information Processing Systems, 2015, pp. 505–513.',
        '[3] G. Abdelrahman, Q. Wang, and B. Nunes, “Knowledge Tracing: A survey,” ACM Computing Surveys, vol. 55, no. 11, pp. 1–37, 2023.',
        '[4] Z. Liu, Q. Liu, J. Chen, S. Huang, and W. Luo, “SimpleKT: A simple but tough-to-beat baseline for Knowledge Tracing,” in The Eleventh International Conference on Learning Representations, 2023.',
        '[5] Z. Liu, Q. Liu, J. Chen, S. Huang, J. Tang, and W. Luo, “pyKT: A python library to benchmark deep learning based Knowledge Tracing models,” in Thirty-seventh Conference on Neural Information Processing Systems Datasets and Benchmarks Track, 2022.',
        '[6] G. Rasch, Probabilistic Models for Some Intelligence and Attainment Tests. Danish Institute for Educational Research, 1960.',
        '[7] A. Ghosh, N. Heffernan, and A. S. Lan, “Context-aware attentive Knowledge Tracing,” in Proc. ACM SIGKDD Conf. Knowledge Discovery and Data Mining, 2020, pp. 2330–2339.',
        '[8] H. Nakagawa, Y. Iwasawa, and Y. Matsuo, “Graph-based Knowledge Tracing: Modeling student proficiency using graph neural network,” in IEEE/WIC/ACM International Conference on Web Intelligence, 2019, pp. 156–163.',
        '[9] W. Lee, J. Chun, Y. Lee, K. Park, and D. Choi, “Contrastive learning for Knowledge Tracing,” in Proc. ACM Web Conference, 2022, pp. 2330–2338.',
        '[10] R. Pelánek, “Metrics for evaluation of student models,” Journal of Educational Data Mining, vol. 7, no. 2, pp. 1–19, 2015.',
        '[11] C. Guo, G. Pleiss, Y. Sun, and K. Q. Weinberger, “On calibration of modern neural networks,” in Proc. International Conference on Machine Learning, 2017, pp. 1321–1330.',
        '[12] M. P. Naeini, G. Cooper, and M. Hauskrecht, “Obtaining well calibrated probabilities using Bayesian binning,” in Proc. AAAI Conference on Artificial Intelligence, vol. 29, no. 1, 2015.',
        '[13] G. W. Brier, “Verification of forecasts expressed in terms of probability,” Monthly Weather Review, vol. 78, no. 1, pp. 1–3, 1950.',
        '[14] M. H. DeGroot and S. E. Fienberg, “The comparison and evaluation of forecasters,” The Statistician, vol. 32, no. 1/2, pp. 12–22, 1983.',
        '[15] X. Yan, C. Tang, and A. Shimada, “Recovering stranded discrimination in Knowledge Tracing: Per-item bias correction via empirical-Bayes shrinkage,” arXiv:2606.14123, 2026.',
        '[16] S. Huang, Z. Liu, X. Zhao, W. Luo, and J. Weng, “Towards robust Knowledge Tracing models via k-sparse attention,” in Proc. 46th International ACM SIGIR Conference, 2023, pp. 2441–2445.',
        '[17] I. Bhattacharjee and C. Wayllace, “Cold start problem: An experimental study of Knowledge Tracing models with new students,” in Artificial Intelligence in Education (AIED 2025), LNCS vol. 15880, Springer, 2025, pp. 425–432.',
        '[18] ASSISTmentsData, “ASSISTments 2012–2013 school data with affect,” ASSISTments Public Datasets, 2012. [Online]. Available: https://sites.google.com/site/assistmentsdata/datasets/2012-13-school-data-with-affect',
        '[19] Junyi Academy, “Junyi Academy online learning activity dataset,” Kaggle, 2019. [Online]. Available: https://www.kaggle.com/datasets/junyiacademy/learning-activity-public-dataset-by-junyi-academy',
        '[20] Z. Liu, Q. Liu, T. Guo, J. Chen, S. Huang, X. Zhao, J. Tang, W. Luo, and J. Weng, “XES3G5M: A Knowledge Tracing benchmark dataset with auxiliary information,” in Thirty-seventh Conference on Neural Information Processing Systems Datasets and Benchmarks Track, 2023.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(11)
        pf.left_indent = Cm(0.5)
        pf.first_line_indent = Cm(-0.5)
        add_text(p, ref, size=8)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    COPY_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    shutil.copy2(OUT_PATH, COPY_PATH)
    print(f"Wrote {OUT_PATH}")
    print(f"Copied {COPY_PATH}")
    print(f"Figure embedded: {fig is not None}")
    return OUT_PATH


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise

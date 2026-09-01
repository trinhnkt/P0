from pathlib import Path
from docx import Document

p = Path(r"C:\TRINH\Sparse-Concept and Calibration\sparse-calibration-kt\IJIET_SUBMISSION\source\main_ijiet_step06.docx")
d = Document(str(p))
out = Path(r"C:\TRINH\Sparse-Concept and Calibration\sparse-calibration-kt\IJIET_SUBMISSION\audit\step06_methods.txt")
lines = []
capture = False
for para in d.paragraphs:
    st = para.style.name
    t = para.text.strip()
    if st == "Heading 1" and "Materials" in t:
        capture = True
        lines.append(f"[{st}] {t}")
        continue
    if capture and st == "Heading 1":
        lines.append(f"STOP [{st}] {t}")
        break
    if capture:
        lines.append(f"[{st}] {t[:800]}")
# tables in methods: first table is table 1
if d.tables:
    t0 = d.tables[0]
    lines.append("\n=== TABLE 1 ===")
    for row in t0.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        lines.append(" | ".join(cells))
out.write_text("\n".join(lines), encoding="utf-8")
print("wrote", out, "nlines", len(lines))

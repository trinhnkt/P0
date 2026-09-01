"""Dump headings, Literature Review, and references from step05."""
from pathlib import Path
from docx import Document

p = Path(r"C:\TRINH\Sparse-Concept and Calibration\sparse-calibration-kt\IJIET_SUBMISSION\source\main_ijiet_step05.docx")
d = Document(str(p))
out = Path(r"C:\TRINH\Sparse-Concept and Calibration\sparse-calibration-kt\IJIET_SUBMISSION\audit\step05_outline.txt")
lines = []
for i, para in enumerate(d.paragraphs, 1):
    st = para.style.name
    t = para.text.strip().replace("\r", " ").replace("\n", " ")
    if st.startswith("Heading") or st in ("Reference Head",) or t.startswith("["):
        lines.append(f"{i:03d} [{st}] {t[:200]}")
    elif "Literature" in t or t.startswith("Classical") or t.startswith("Educational"):
        lines.append(f"{i:03d} [{st}] {t}")

# full lit window
lines.append("\n=== LITERATURE REVIEW FULL ===")
capture = False
for para in d.paragraphs:
    st = para.style.name
    t = para.text.strip()
    if st == "Heading 1" and "Literature" in t:
        capture = True
        lines.append(f"[{st}] {t}")
        continue
    if capture and st == "Heading 1":
        lines.append(f"STOP [{st}] {t}")
        break
    if capture:
        lines.append(f"[{st}] {t}")

lines.append("\n=== REFERENCES ===")
for para in d.paragraphs:
    t = para.text.strip()
    if t.startswith("[") and t[1:3].rstrip("]").isdigit() or (
        t.startswith("[") and "]" in t[:5]
    ):
        lines.append(t[:400])

out.write_text("\n".join(lines), encoding="utf-8")
print(out.read_text(encoding="utf-8"))

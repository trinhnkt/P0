"""Dump reference paragraphs and Heading 2 list strings from step05."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

p = Path(r"C:\TRINH\Sparse-Concept and Calibration\sparse-calibration-kt\IJIET_SUBMISSION\source\main_ijiet_step05.docx")
d = Document(str(p))
out = Path(r"C:\TRINH\Sparse-Concept and Calibration\sparse-calibration-kt\IJIET_SUBMISSION\audit\step05_refs.txt")
lines = []
for i, para in enumerate(d.paragraphs, 1):
    st = para.style.name
    t = para.text.strip()
    if st in ("References", "Reference Head") or "Corbett" in t or "Piech" in t:
        lines.append(f"{i:03d} [{st}] {t}")
    if st == "Heading 2":
        lines.append(f"{i:03d} H2 {t}")
# last 40 paragraphs
lines.append("\n=== LAST 40 PARAS ===")
paras = list(d.paragraphs)
for i, para in enumerate(paras[-40:], start=len(paras) - 39):
    lines.append(f"{i:03d} [{para.style.name}] {para.text.strip()[:350]}")
out.write_text("\n".join(lines), encoding="utf-8")
print(out.read_text(encoding="utf-8"))

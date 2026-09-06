#!/usr/bin/env python3
"""A31 text patch via python-docx (Word COM license unavailable)."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

HERE = Path(__file__).resolve().parent
ABS_OLD = (
    "The primary check is 4/4 unique partitions (mean ΔFAR 0.047; five runs, "
    "two sharing a split); the seed-42 95% CI is [0.006, 0.138]."
)
ABS_NEW = (
    "The primary check is 4/4 unique learner partitions (partition-level mean "
    "ΔFAR 0.056, range 0.015–0.087); the five training runs that underlie those "
    "partitions have mean 0.047 (sd 0.033) because seeds 2025 and 2026 share one "
    "split. The seed-42 95% CI is [0.006, 0.138]."
)
CAP_OLD = (
    "Table 6. Gate robustness at τ=0.7 on ASSISTments 2012 (primary unit: four "
    "unique learner partitions). Five training runs (seeds 42, 2024, 2025, 2026, "
    "2027) across those partitions (2025 and 2026 share a split). Partition-level "
    "ΔFAR averages seeds 2025 and 2026 first; T-KT mean 0.056, range 0.015–0.087. "
    "Mean N, Nadvance, and Nincorrect are sparse-stratum denominators. "
    "GKT/CL4KT remain seed 42 only."
)
CAP_NEW = (
    "Table 6. Gate robustness at τ=0.7 on ASSISTments 2012 (primary unit: four "
    "unique learner partitions). T-KT Mean ΔFAR is the partition-level mean 0.056 "
    "(the SD column is the partition range 0.015–0.087) after averaging seeds "
    "2025 and 2026 first. The five training runs have mean 0.047 (sd 0.033) "
    "because those two seeds share one split. DKT remains a five-run summary. "
    "Mean N, Nadvance, and Nincorrect are sparse-stratum denominators. "
    "GKT/CL4KT remain seed 42 only."
)


def replace_in_runs(p, old: str, new: str) -> bool:
    if old not in p.text:
        return False
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new, 1)
            return True
    raise SystemExit("old string split across runs")


def patch(path: Path) -> None:
    d = Document(str(path))
    n_abs = n_cap = n_tbl = 0
    for p in d.paragraphs:
        if replace_in_runs(p, ABS_OLD, ABS_NEW):
            n_abs += 1
            # keep Abstract— label if we wiped run 0
            if p.text.startswith("Knowledge Tracing") and not p.text.startswith(
                "Abstract—"
            ):
                raise SystemExit("lost Abstract— label")
        if replace_in_runs(p, CAP_OLD, CAP_NEW):
            n_cap += 1
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 4:
                continue
            if cells[0] != "T-KT" or cells[1] != "0.047":
                continue
            if "4/4 unique partitions" not in cells[3]:
                continue
            row.cells[1].text = "0.056"
            row.cells[2].text = "0.015–0.087"
            row.cells[3].text = "4/4 unique partitions"
            n_tbl += 1
    if n_abs != 1 or n_cap != 1 or n_tbl != 1:
        raise SystemExit(f"{path.name}: abs={n_abs} cap={n_cap} tbl={n_tbl}")
    if ABS_OLD in d.paragraphs[0].text:
        raise SystemExit("old abstract still in p0")
    d.save(str(path))
    print("patched", path.name)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    for name in ("main_ijiet_full.docx", "main_ijiet_blind.docx"):
        patch(HERE / "manuscript" / name)


if __name__ == "__main__":
    main()

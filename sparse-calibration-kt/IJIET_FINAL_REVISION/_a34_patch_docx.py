#!/usr/bin/env python3
"""A34: drop GKT/CL4KT CIs from the Table 5 note; drop orphan Table 6 caption clause."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

HERE = Path(__file__).resolve().parent
NOTE_OLD = "; GKT [\u22120.054, 0.092]; CL4KT [\u22120.018, 0.142]."
NOTE_NEW = "."
CAP_TAIL = " GKT/CL4KT remain seed 42 only."
IVE_GKT = "[\u22120.054, 0.092]"
IVE_CL4 = "[\u22120.018, 0.142]"


def replace_in_runs(p, old: str, new: str) -> bool:
    if old not in p.text:
        return False
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new, 1)
            return True
    raise SystemExit(f"old string split across runs: {old[:40]!r}")


def patch(path: Path) -> None:
    d = Document(str(path))
    n_note = n_cap = 0
    for p in d.paragraphs:
        if replace_in_runs(p, NOTE_OLD, NOTE_NEW):
            n_note += 1
        if replace_in_runs(p, CAP_TAIL, ""):
            n_cap += 1
    full = "\n".join(p.text for p in d.paragraphs)
    if n_note != 1 or n_cap != 1:
        raise SystemExit(f"{path.name}: note={n_note} cap={n_cap}")
    if "GKT/CL4KT remain seed 42 only" in full:
        raise SystemExit(f"{path.name}: Table 6 tail still present")
    if NOTE_OLD in full:
        raise SystemExit(f"{path.name}: GKT/CL4KT still on Table 5 note")
    if IVE_GKT not in full or IVE_CL4 not in full:
        raise SystemExit(f"{path.name}: IV.E CIs missing")
    if "not shown in Table 5" not in full:
        raise SystemExit(f"{path.name}: IV.E Table 5 xref missing")
    d.save(str(path))
    print("patched", path.name)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    for name in ("main_ijiet_full.docx", "main_ijiet_blind.docx"):
        patch(HERE / "manuscript" / name)


if __name__ == "__main__":
    main()

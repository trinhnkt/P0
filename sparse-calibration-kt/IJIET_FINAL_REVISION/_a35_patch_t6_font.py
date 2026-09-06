#!/usr/bin/env python3
"""A35: restore Table 6 T-KT Mean/SD/Runs cells to 7 pt TNR (A31 cell.text reset)."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

HERE = Path(__file__).resolve().parent
TARGETS = {
    1: "0.056",
    2: "0.015–0.087",
    3: "4/4 unique partitions",
}


def style_cell(cell, expected: str) -> None:
    text = cell.text.strip().replace("\n", " ")
    if text != expected and text.replace("  ", " ") != expected:
        raise SystemExit(f"cell text {text!r} != {expected!r}")
    n = 0
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(7)
            r.font.name = "Times New Roman"
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:eastAsia"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")
            n += 1
    if n < 1:
        raise SystemExit(f"no runs in {expected!r}")


def patch(path: Path) -> None:
    d = Document(str(path))
    hits = 0
    for tbl in d.tables:
        if len(tbl.rows) < 2 or len(tbl.columns) < 4:
            continue
        row = tbl.rows[1]
        if row.cells[0].text.strip() != "T-KT":
            continue
        if "0.056" not in row.cells[1].text:
            continue
        for idx, expected in TARGETS.items():
            style_cell(row.cells[idx], expected)
        hits += 1
    if hits != 1:
        raise SystemExit(f"{path.name}: T-KT Table 6 hits={hits}")
    d.save(str(path))
    print("patched", path.name)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    for name in ("main_ijiet_full.docx", "main_ijiet_blind.docx"):
        patch(HERE / "manuscript" / name)


if __name__ == "__main__":
    main()
